import streamlit as st
from docx import Document
from openai import OpenAI
import json
import re

client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

st.title("AI Генератор учебных материалов")

teacher = st.text_input("ФИО преподавателя")
subject = st.text_input("Название предмета")
topic = st.text_input("Тема урока")

group = st.text_input("Группа")

course = st.selectbox(
    "Курс",
    ["1","2","3","4"]
)

lesson_duration = st.selectbox(
    "Длительность урока",
    ["60 минут","80 минут"]
)

date = st.date_input("Дата урока")

language = st.radio(
    "Выберите язык",
    ["Русский","Қазақша"]
)

lesson_type_choice = st.selectbox(
    "Тип урока",
    [
        "Усвоения новых знаний / Жаңа білім беру сабағы",
        "Закрепления знаний, умений и навыков / Бекіту сабақ",
        "Повторения и актуализации знания и умений / Қайталау сабақ",
        "Обобщение и систематизации знаний и умений / Жалпылау сабақ",
        "Подведение итога знаний и умений / Қорытындылау сабақ",
        "Комбинированный урок / Аралас сабақ"
    ]
)

# распределение времени

if lesson_duration == "60 минут":

    time1="5 мин"
    time2="15 мин"
    time3="10 мин"
    time4="15 мин"
    time5="5 мин"
    time6="5 мин"
    time7="5 мин"

else:

    time1="5 мин"
    time2="20 мин"
    time3="15 мин"
    time4="20 мин"
    time5="10 мин"
    time6="5 мин"
    time7="5 мин"


# локализация

if language == "Русский":

    lang_instruction="Напиши текст на русском языке."

    org_steps="""1. Приветствие обучающихся
2. Проверка присутствующих
3. Ознакомление с темой и целью занятия"""

    lesson_type=lesson_type_choice.split("/")[0].strip()

    presentation_resource="Презентация к уроку"
    presentation_demo="Демонстрация на презентации."
    peer_assessment="Взаимооценивание студентов."
    assessment_sheet="Лист оценивания."

else:

    lang_instruction="Жауаптың барлығын қазақ тілінде жаз."

    org_steps="""1. Оқушыларды қарсы алу
2. Қатысушыларды тексеру
3. Сабақтың тақырыбы мен мақсатымен танысу"""

    lesson_type=lesson_type_choice.split("/")[1].strip()

    presentation_resource="Сабаққа презентация"
    presentation_demo="Презентация демонстрациясы."
    peer_assessment="Студенттерді өзара бағалау."
    assessment_sheet="Бағалау парағы"


# ===============================
# ПЛАН УРОКА
# ===============================

if st.button("📄 Сгенерировать план урока"):

    prompt=f"""
Создай подробный план учебного занятия для колледжа.

Предмет: {subject}
Тема: {topic}

Продолжительность урока: {lesson_duration}

Используй таксономию Блума.

Структура:

1. Организационный этап (Знание)
2. Формирование новых знаний (Понимание)
3. Формирование умений и навыков (Применение)
4. Анализ
5. Подведение итогов
6. Рефлексия
7. Домашнее задание

Язык: {language}

Верни JSON:

{{
"goal":"",
"tasks":"",
"results":"",
"resources":"",
"theory":"",
"practice":"",
"practice_result":"",
"reflection":"",
"homework":""
}}
"""

    response=client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role":"user","content":prompt}]
    )

    text=response.choices[0].message.content

    json_text=re.search(r'\{.*\}',text,re.DOTALL)

    data=json.loads(json_text.group())

    doc=Document("template.docx")

    replacements={

        "{teacher}":teacher,
        "{subject}":subject,
        "{topic}":topic,
        "{group}":group,
        "{course}":course,
        "{date}":date.strftime("%d.%m.%Y"),

        "{lesson_type}":lesson_type,

        "{time1}":time1,
        "{time2}":time2,
        "{time3}":time3,
        "{time4}":time4,
        "{time5}":time5,
        "{time6}":time6,
        "{time7}":time7,

        "{goal}":data.get("goal",""),
        "{tasks}":data.get("tasks",""),
        "{results}":data.get("results",""),
        "{resources}":data.get("resources",""),

        "{org_steps}":org_steps,

        "{theory}":data.get("theory",""),
        "{practice}":data.get("practice",""),
        "{practice_result}":data.get("practice_result",""),

        "{reflection}":data.get("reflection",""),
        "{homework}":data.get("homework",""),

        "{presentation_resource}":presentation_resource,
        "{presentation_demo}":presentation_demo,
        "{peer_assessment}":peer_assessment,
        "{assessment_sheet}":assessment_sheet
    }

    # замена текста

    for p in doc.paragraphs:
        for k,v in replacements.items():
            if k in p.text:
                p.text = p.text.replace(k, str(v))

    # замена в таблицах

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k,v in replacements.items():
                        if k in p.text:
                            p.text = p.text.replace(k, str(v))

    file="lesson_plan.docx"
    doc.save(file)

    with open(file,"rb") as f:
        st.download_button("Скачать план урока",f,file_name=file)


# ===============================
# ЛЕКЦИЯ
# ===============================

if st.button("📘 Сгенерировать лекцию"):

    prompt=f"""
Напиши лекцию для колледжа.

Тема: {topic}

Размер: примерно 2 страницы A4.

Структура:
введение
основная часть
заключение

Язык: {language}
"""

    response=client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role":"user","content":prompt}]
    )

    lecture=response.choices[0].message.content

    doc=Document()
    doc.add_heading(topic,0)

    for p in lecture.split("\n"):
        doc.add_paragraph(p)

    file="lecture.docx"
    doc.save(file)

    with open(file,"rb") as f:
        st.download_button("Скачать лекцию",f,file_name=file)


# ===============================
# ПРАКТИЧЕСКАЯ
# ===============================

if st.button("🧪 Сгенерировать практическую работу"):

    prompt=f"""
Создай практическую работу.

Тема: {topic}

3 задания.

Язык: {language}
"""

    response=client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role":"user","content":prompt}]
    )

    text=response.choices[0].message.content

    doc=Document()
    doc.add_heading("Практическая работа",0)

    for p in text.split("\n"):
        doc.add_paragraph(p)

    file="practice.docx"
    doc.save(file)

    with open(file,"rb") as f:
        st.download_button("Скачать практическую работу",f,file_name=file)


# ===============================
# ТЕСТ
# ===============================

if st.button("📝 Сгенерировать тест"):

    prompt=f"""
Создай тест из 10 вопросов.

Тема: {topic}

4 варианта ответа.

Язык: {language}
"""

    response=client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role":"user","content":prompt}]
    )

    text=response.choices[0].message.content

    doc=Document()
    doc.add_heading("Тест",0)

    for p in text.split("\n"):
        doc.add_paragraph(p)

    file="test.docx"
    doc.save(file)

    with open(file,"rb") as f:
        st.download_button("Скачать тест",f,file_name=file)