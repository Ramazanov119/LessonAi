import json
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt

from config.colleges import get_college_config
from services.document_style import (
    add_page_number,
    configure_document,
    normalize_document_fonts,
    normalize_paragraph,
    set_run_font,
)


LESSON_TEMPLATE = get_college_config("ETEC")["template"]


def _add_block(doc, title, body):
    doc.add_heading(title, level=2)
    para = doc.add_paragraph()
    run = para.add_run(str(body))
    set_run_font(run)
    normalize_paragraph(para)


def _get_logo_path(college):
    logo_path = get_college_config(college)["logo"]
    return logo_path if logo_path.exists() else None


def _add_title_page(
    doc,
    *,
    college=None,
    subject=None,
    title=None,
    teacher=None,
    group=None,
    course=None,
    lesson_date=None,
):
    logo_path = _get_logo_path(college)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_after = Pt(10)

    if logo_path:
        try:
            title_para.add_run().add_picture(
                str(logo_path),
                width=Inches(2.2),
            )
        except Exception:
            pass

    meta = [
        ("Преподаватель", teacher or ""),
        (
            "Дата",
            lesson_date.strftime("%d.%m.%Y")
            if lesson_date
            else "",
        ),
    ]

    for label, value in meta:
        if not value:
            continue

        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line.paragraph_format.space_after = Pt(3)

        label_run = line.add_run(f"{label}: ")
        label_run.bold = True
        set_run_font(label_run)

        value_run = line.add_run(str(value))
        set_run_font(value_run)

    subject_text = str(subject or "Предмет")

    subject_para = doc.add_paragraph()
    subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subject_para.paragraph_format.space_before = Pt(48)
    subject_para.paragraph_format.space_after = Pt(8)

    subject_run = subject_para.add_run(subject_text)
    set_run_font(subject_run, Pt(16), bold=True)

    type_para = doc.add_paragraph()
    type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_para.paragraph_format.space_after = Pt(8)

    type_run = type_para.add_run("ЛЕКЦИЯ")
    set_run_font(type_run, Pt(16), bold=True)

    lecture_title = str(title or "Лекция")

    lecture_para = doc.add_paragraph()
    lecture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lecture_para.paragraph_format.space_after = Pt(0)

    lecture_run = lecture_para.add_run(lecture_title)
    set_run_font(lecture_run, Pt(16), bold=True)

    doc.add_page_break()


def _add_list_item(doc, text, level=0):
    style_name = "List Bullet" if level == 0 else "List Bullet 2"

    if style_name not in doc.styles:
        style_name = "List Bullet"

    paragraph = doc.add_paragraph(style=style_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)

    paragraph.add_run(str(text))

    for run in paragraph.runs:
        set_run_font(run)


def _add_info_box(doc, title, content):
    doc.add_heading(title, level=2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.style = "Intense Quote"

    run = p.add_run(str(content))
    set_run_font(run)

    normalize_paragraph(p)


def _strip_json_block(text):
    if not isinstance(text, str):
        return text

    cleaned = text.strip()

    if cleaned.startswith("```"):
        cleaned = (
            cleaned
            .split("\n", 1)[-1]
            .rsplit("```", 1)[0]
            .strip()
        )

    return cleaned


def _lesson_fields(content):
    if isinstance(content, dict):
        return content

    try:
        parsed = json.loads(_strip_json_block(content))
    except (TypeError, json.JSONDecodeError):
        return {"theory": content}

    return (
        parsed
        if isinstance(parsed, dict)
        else {"theory": content}
    )


def _replace_placeholders(paragraph, fields):
    text = paragraph.text

    for key, value in fields.items():
        text = text.replace(
            "{" + key + "}",
            "" if value is None else str(value),
        )

    if text != paragraph.text:
        paragraph.text = text

    if not paragraph.runs:
        paragraph.add_run()
    for run in paragraph.runs:
        set_run_font(run, size=Pt(11))


def _fill_document(doc, fields):
    for paragraph in doc.paragraphs:
        _replace_placeholders(paragraph, fields)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_placeholders(paragraph, fields)


def _meta_template_fields(content):
    """
    Prepare the fields used by the META lesson-plan template.

    The original lesson fields are not removed. These fields are
    additionally prepared for the template so that META-specific
    placeholders can be filled without breaking the existing ETEC
    generation flow.
    """
    payload = _lesson_fields(content)

    return {
        # Организационный этап
        "time1": payload.get("time1", ""),
        "org_steps": payload.get("org_steps", ""),

        # Актуализация
        "time2": payload.get("time2", ""),
        "actualization": payload.get("actualization", ""),
        "actualization_resource": payload.get(
            "actualization_resource",
            "",
        ),

        # Теория
        "time3": payload.get("time3", ""),
        "theory": payload.get("theory", ""),
        "presentation_demo": payload.get(
            "presentation_demo",
            "",
        ),

        # Практика
        "time4": payload.get("time4", ""),
        "practice": payload.get("practice", ""),
        "practice_result": payload.get(
            "practice_result",
            "",
        ),

        # Взаимооценивание
        "time5": payload.get("time5", ""),
        "peer_assessment": payload.get(
            "peer_assessment",
            "",
        ),

        # Рефлексия и оценивание
        "time6": payload.get("time6", ""),
        "reflection": payload.get("reflection", ""),
        "assessment": payload.get("assessment", ""),
        "assessment_sheet": payload.get(
            "assessment_sheet",
            "",
        ),

        # Домашнее задание
        "homework": payload.get("homework", ""),

        # Дополнительные поля шаблона
        "date": payload.get("date", ""),
        "evaluation_criteria": payload.get(
            "evaluation_criteria",
            "",
        ),
        "explanatory_illustrative_method": payload.get(
            "explanatory_illustrative_method",
            payload.get("presentation_demo", ""),
        ),
        "practical_task": payload.get(
            "practical_task",
            payload.get("practice", ""),
        ),
        "interdisciplinary_connections": payload.get(
            "interdisciplinary_connections",
            payload.get("resources", ""),
        ),
        "resources": payload.get(
            "resources",
            "",
        ),
    }


def build_doc(
    title,
    content,
    schema=None,
    visual=None,
    template_path=LESSON_TEMPLATE,
    fields=None,
):
    if title == "План урока" and Path(template_path).exists():
        doc = Document(template_path)

        # Основные поля, которые уже существовали в проекте.
        template_fields = _lesson_fields(content)

        # Дополнительные поля META-шаблона.
        meta_fields = _meta_template_fields(content)

        # Объединяем оба набора.
        #
        # template_fields сохраняет обратную совместимость
        # с существующим ETEC-шаблоном.
        #
        # meta_fields добавляет специальные поля META.
        #
        # fields имеет самый высокий приоритет, если какие-либо
        # значения передаются непосредственно в build_doc().
        all_fields = {
            **template_fields,
            **meta_fields,
            **(fields or {}),
        }

        _fill_document(doc, all_fields)
        configure_document(doc)
        normalize_document_fonts(doc)

    else:
        doc = Document()

        title_style = doc.styles["Normal"]
        title_style.font.name = "Times New Roman"
        title_style.font.size = Pt(12)

        doc.add_heading(title, 1)
        doc.add_paragraph(content)

    if schema and title != "План урока":
        doc.add_heading("Схема", 2)
        doc.add_paragraph(schema)

    if visual and title != "План урока":
        doc.add_heading("Иллюстрация", 2)
        doc.add_paragraph(visual)

    file = BytesIO()
    doc.save(file)
    file.seek(0)

    return file


def build_lecture_docx(
    lecture_text,
    title="Лекция",
    *,
    college=None,
    subject=None,
    teacher=None,
    group=None,
    course=None,
    lesson_date=None,
):
    payload = _lesson_fields(lecture_text)

    doc = Document()
    configure_document(doc)

    _add_title_page(
        doc,
        college=college,
        subject=subject,
        title=title,
        teacher=teacher,
        group=group,
        course=course,
        lesson_date=lesson_date,
    )

    if payload.get("relevance"):
        doc.add_heading("Актуальность", level=1)

        paragraph = doc.add_paragraph(
            payload["relevance"]
        )
        normalize_paragraph(paragraph)

    if payload.get("learning_goals"):
        doc.add_heading("Цели обучения", level=1)

        for item in payload["learning_goals"]:
            _add_list_item(doc, str(item))

    if payload.get("key_terms"):
        doc.add_heading("Ключевые понятия", level=1)

        for term in payload["key_terms"]:
            doc.add_heading(
                str(term.get("term", "Понятие")),
                level=2,
            )

            if term.get("definition"):
                paragraph = doc.add_paragraph(
                    str(term.get("definition", ""))
                )
                normalize_paragraph(paragraph)

    if payload.get("main_material"):
        doc.add_heading(
            "Основной теоретический материал",
            level=1,
        )

        for section in payload["main_material"]:
            doc.add_heading(
                str(section.get("title", "Материал")),
                level=2,
            )

            if section.get("content"):
                paragraph = doc.add_paragraph(
                    str(section.get("content", ""))
                )
                normalize_paragraph(paragraph)

            if section.get("examples"):
                doc.add_heading("Примеры", level=3)

                for example in section.get("examples", []):
                    _add_list_item(doc, str(example))

    if payload.get("practice_examples"):
        doc.add_heading(
            "Практические примеры",
            level=1,
        )

        for item in payload["practice_examples"]:
            _add_list_item(doc, str(item))

    if payload.get("case_study"):
        case = payload["case_study"]

        doc.add_heading(
            "Профессиональная ситуация / кейс",
            level=1,
        )

        _add_block(
            doc,
            "Сценарий",
            case.get("scenario", ""),
        )

        _add_block(
            doc,
            "Анализ",
            case.get("analysis", ""),
        )

        _add_block(
            doc,
            "Решение",
            case.get("solution", ""),
        )

    if payload.get("common_errors"):
        doc.add_heading(
            "Типичные ошибки",
            level=1,
        )

        for item in payload["common_errors"]:
            _add_block(
                doc,
                "Ошибка",
                item.get("error", ""),
            )

            _add_block(
                doc,
                "Почему возникает",
                item.get("why", ""),
            )

            _add_block(
                doc,
                "Как исправить",
                item.get("fix", ""),
            )

    if payload.get("summary"):
        doc.add_heading("Резюме", level=1)

        for item in payload["summary"]:
            _add_list_item(doc, str(item))

    if payload.get("review_questions"):
        doc.add_heading(
            "Вопросы для закрепления",
            level=1,
        )

        for index, item in enumerate(
            payload["review_questions"],
            start=1,
        ):
            paragraph = doc.add_paragraph(
                f"{index}. {item}"
            )
            normalize_paragraph(paragraph)

    if payload.get("preparation_for_practice"):
        doc.add_heading(
            "Подготовка к практическому занятию",
            level=1,
        )

        paragraph = doc.add_paragraph(
            str(payload["preparation_for_practice"])
        )
        normalize_paragraph(paragraph)

    for section in doc.sections:
        para = section.footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.text = ""

        footer_label = para.add_run("Страница ")
        set_run_font(
            footer_label,
            size=Pt(12),
        )

        add_page_number(para)

    normalize_document_fonts(doc)

    file = BytesIO()
    doc.save(file)
    file.seek(0)

    return file


def build_practice_docx(
    practice_text,
    title="Практическое занятие",
    *,
    college=None,
    subject=None,
    teacher=None,
    lesson_date=None,
):
    """Build a styled DOCX from the stored structured practice content."""

    payload = _lesson_fields(practice_text)

    doc = Document()
    configure_document(doc)

    _add_title_page(
        doc,
        college=college,
        subject=subject,
        title=title,
        teacher=teacher,
        lesson_date=lesson_date,
    )

    sections = (
        ("Цель", "objective"),
        ("Краткая теория", "brief_theory"),
        ("Профессиональная ситуация", "professional_context"),
        ("Практическое задание", "main_task"),
        ("Ожидаемый результат", "expected_result"),
        ("Заключение", "conclusion"),
    )

    doc.add_heading(
        str(payload.get("title", title)),
        level=1,
    )

    for heading, key in sections:
        if payload.get(key):
            doc.add_heading(heading, level=1)

            paragraph = doc.add_paragraph(
                str(payload[key])
            )
            normalize_paragraph(paragraph)

    list_sections = (
        ("Ожидаемые результаты", "learning_outcomes"),
        ("Необходимые инструменты", "required_tools"),
        ("Техника безопасности", "safety_notes"),
        ("Пошаговая инструкция", "task_steps"),
        ("Индивидуальные варианты", "individual_variants"),
        ("Контрольные вопросы", "control_questions"),
        ("Рефлексия", "reflection"),
    )

    for heading, key in list_sections:
        values = payload.get(key) or []

        if values:
            doc.add_heading(heading, level=1)

            for value in values:
                _add_list_item(doc, str(value))

    criteria = payload.get("evaluation_criteria") or []

    if criteria:
        doc.add_heading(
            "Критерии оценки",
            level=1,
        )

        for item in criteria:
            if isinstance(item, dict):
                _add_list_item(
                    doc,
                    f"{item.get('criterion', '')} — "
                    f"{item.get('weight', '')}",
                )
            else:
                _add_list_item(doc, str(item))

    time_allocation = payload.get(
        "time_allocation"
    ) or []

    if time_allocation:
        doc.add_heading(
            "Распределение времени — 70 минут",
            level=1,
        )

        for item in time_allocation:
            if isinstance(item, dict):
                _add_list_item(
                    doc,
                    f"{item.get('stage', '')}: "
                    f"{item.get('minutes', '')} мин.",
                )
            else:
                _add_list_item(doc, str(item))

    normalize_document_fonts(doc)

    file = BytesIO()
    doc.save(file)
    file.seek(0)

    return file