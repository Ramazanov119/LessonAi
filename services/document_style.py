from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(11)
TABLE_SIZE = Pt(11)
CAPTION_SIZE = Pt(11)
FIRST_LINE_INDENT = Cm(1.25)


def set_run_font(run, size=BODY_SIZE, bold=None, italic=None):
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = size
    run.font.color.rgb = None
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = None
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_settings = {
        "Heading 1": (Pt(16), WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 2": (Pt(14), WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": (Pt(14), WD_ALIGN_PARAGRAPH.LEFT),
    }
    for style_name, (size, alignment) in heading_settings.items():
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = size
        style.font.bold = True
        style.font.color.rgb = None
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.first_line_indent = Cm(0)

    for style_name in ("List Bullet", "List Number", "Intense Quote"):
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = BODY_SIZE
        style.font.color.rgb = None
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(3)

    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)
        section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def normalize_paragraph(paragraph, *, size=BODY_SIZE, first_line_indent=FIRST_LINE_INDENT):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = first_line_indent
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    if not paragraph.runs:
        paragraph.add_run()
    for run in paragraph.runs:
        set_run_font(run, size=size)


def normalize_document_fonts(doc):
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            size = doc.styles[paragraph.style.name].font.size or BODY_SIZE
        else:
            size = BODY_SIZE
        for run in paragraph.runs:
            set_run_font(run, size=run.font.size or size)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    for run in paragraph.runs:
                        set_run_font(run, size=TABLE_SIZE)


def add_page_number(paragraph):
    run = paragraph.add_run()
    set_run_font(run, size=Pt(12))
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, result, end))


def add_caption(doc, text, *, table=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=CAPTION_SIZE, italic=not table)
    return paragraph
